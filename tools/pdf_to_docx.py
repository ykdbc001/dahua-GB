# -*- coding: utf-8 -*-
"""将目录内消防协议 PDF 转为 Word（纯文本抽取，版式/图片以原 PDF 为准）。"""
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def main():
    base = Path(__file__).resolve().parent.parent
    pdfs = sorted(base.glob("*.pdf"))
    if not pdfs:
        raise SystemExit(f"未在目录找到 PDF: {base}")

    pdf_path = pdfs[0]
    if len(pdfs) > 1:
        # 优先文件名含「国标」或「消防终端」
        for p in pdfs:
            n = p.name
            if "国标" in n or "消防终端" in n or "消防产品线" in n:
                pdf_path = p
                break

    out_path = pdf_path.with_suffix(".docx")
    print("OK: reading PDF pages...")

    reader = PdfReader(str(pdf_path))
    doc = Document()
    doc.add_heading(pdf_path.stem, level=0)
    doc.add_paragraph(f"由 PDF 自动抽取文本生成，共 {len(reader.pages)} 页。表格与排版可能与原 PDF 不一致。")

    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        doc.add_heading(f"第 {i} 页", level=2)
        for line in text.splitlines():
            line = line.rstrip()
            if line:
                doc.add_paragraph(line)
            else:
                doc.add_paragraph("")

    doc.save(str(out_path))
    print("OK: docx saved (same folder as pdf)")


if __name__ == "__main__":
    main()
